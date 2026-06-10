import time
import torch
import gc
import os
from transformers import AutoTokenizer, AutoModelForCausalLM
from datetime import datetime
from docx import Document # Import untuk Word
from docx.shared import Inches

def get_vram_usage():
    if torch.cuda.is_available():
        return torch.cuda.memory_allocated() / (1024**2)
    return 0

QUESTIONS = [
    {"q": "Apa rukun Islam yang ke-4?", "lang": "Indonesia"},
    {"q": "What are the pillars of Iman in Islam?", "lang": "English"},
    {"q": "كم عدد ركعات صلاة الفجر؟", "lang": "Arabic"},
    {"q": "ما معني إهدنا الصراط المستقيم ؟", "lang": "Arabic"},
    {"q": "Explain the concept of Halal in simple words.", "lang": "English"},
    {"q": "Sebutkan 5 sifat Allah untuk berdoa minta rejeki.", "lang": "Indonesia"}
]

def save_to_word(res):
    doc = Document()
    doc.add_heading('Laporan Pengujian SLM Islami', 0)
    
    doc.add_paragraph(f"Waktu Pengujian: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Hardware: MSI Cyborg 15 (NVIDIA RTX 4060)")
    doc.add_paragraph(f"Model: {res['name']}")
    
    doc.add_heading('Tabel Hasil Analisis', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Bahasa'
    hdr_cells[2].text = 'Waktu (s)'
    hdr_cells[3].text = 'Memori (MB)'
    
    for r in res['details']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r['id'])
        row_cells[1].text = r['lang']
        row_cells[2].text = f"{r['time']:.2f}"
        row_cells[3].text = f"{r['mem']:.2f}"
    
    doc.add_heading('Detail Jawaban Model', level=1)
    for r in res['details']:
        p = doc.add_paragraph()
        p.add_run(f"Pertanyaan ({r['lang']}): ").bold = True
        p.add_run(QUESTIONS[r['id']-1]['q'])
        p = doc.add_paragraph()
        p.add_run("Jawaban Model: ").bold = True
        p.add_run(r['answer'])
        doc.add_paragraph("-" * 30)

    filename = "Hasil_Tugas_AI_Islami.docx"
    doc.save(filename)
    print(f"\n[SUKSES] File Word berhasil dibuat: {filename}")

def benchmark_model(model_path, model_name):
    print(f"\nSTART ANALISIS: {model_name}")
    print("-" * 60)
    
    if not os.path.exists(model_path):
        print(f"Error: Folder {model_path} tidak ditemukan!")
        return None

    start_load = time.time()
    try:
        tokenizer = AutoTokenizer.from_pretrained(model_path)
        model = AutoModelForCausalLM.from_pretrained(
            model_path,
            torch_dtype=torch.float16,
            device_map="auto"
        )
    except Exception as e:
        print(f"Gagal memuat model: {e}")
        return None
    
    load_time = time.time() - start_load
    results = []
    
    for i, item in enumerate(QUESTIONS, 1):
        torch.cuda.empty_cache()
        gc.collect()
        mem_before = get_vram_usage()
        start_time = time.time()
        
        inputs = tokenizer(item["q"], return_tensors="pt").to("cuda")
        outputs = model.generate(
            **inputs, 
            max_new_tokens=150, # Lebih panjang sedikit agar jawaban Word lengkap
            do_sample=True, 
            temperature=0.3,
            pad_token_id=tokenizer.eos_token_id
        )
        
        duration = time.time() - start_time
        mem_after = get_vram_usage()
        answer = tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
        
        results.append({
            'id': i, 'lang': item["lang"], 'time': duration,
            'mem': max(0, mem_after - mem_before), 'answer': answer.strip()
        })
        print(f"Selesai Soal {i} | {duration:.2f}s")

    return {'name': model_name, 'load': load_time, 'details': results}

# Eksekusi
print("MEMBERSIHKAN MEMORI...")
torch.cuda.empty_cache()

res = benchmark_model("./qwen-model", "Qwen-0.5B-Islamic")

if res:
    # Tampilkan di Terminal
    print("\n" + "="*80)
    print(f"{'No':<4} {'Bahasa':<12} {'Waktu (s)':<12} {'Memori (MB)':<15}")
    for r in res['details']:
        print(f"{r['id']:<4} {r['lang']:<12} {r['time']:<12.2f} {r['mem']:<15.2f}")
    
    # Simpan ke Word
    save_to_word(res)
    print("="*80)